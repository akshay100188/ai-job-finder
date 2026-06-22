from pathlib import Path

from docx import Document
from docx.shared import Pt

from jobhorizon.master_resume import MasterResume

# bullets_by_source maps "experience:<company>" / "project:<name>" -> list of
# (original_bullet, final_bullet) pairs, in relevance order, for entries that had
# at least one bullet selected during tailoring.
BulletsBySource = dict[str, list[tuple[str, str]]]


def _contact_line(master: MasterResume) -> str:
    parts = [
        master.contact.email,
        master.contact.phone,
        master.contact.location,
        master.contact.linkedin,
        master.contact.website,
    ]
    return " | ".join(p for p in parts if p)


def render_docx(
    master: MasterResume,
    bullets_by_source: BulletsBySource,
    path: Path,
    include_summary: bool = True,
) -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(master.contact.name or "Resume", level=1)
    contact_line = _contact_line(master)
    if contact_line:
        doc.add_paragraph(contact_line)

    if include_summary and master.summary:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(master.summary)

    doc.add_heading("Experience", level=2)
    for exp in master.experience:
        pairs = bullets_by_source.get(f"experience:{exp.company}")
        bullets = [final for _, final in pairs] if pairs else list(exp.bullets)

        header = doc.add_paragraph()
        header.add_run(f"{exp.role}, {exp.company}").bold = True
        meta = f"{exp.start} - {exp.end}" if (exp.start or exp.end) else ""
        if exp.location:
            meta = f"{meta} | {exp.location}" if meta else exp.location
        if meta:
            doc.add_paragraph(meta)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    project_entries = [
        (proj, bullets_by_source[f"project:{proj.name}"])
        for proj in master.projects
        if f"project:{proj.name}" in bullets_by_source
    ]
    if project_entries:
        doc.add_heading("Projects", level=2)
        for proj, pairs in project_entries:
            header = doc.add_paragraph()
            header.add_run(proj.name).bold = True
            for _, final in pairs:
                doc.add_paragraph(final, style="List Bullet")

    if master.skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(master.skills))

    if master.certifications:
        doc.add_heading("Certifications", level=2)
        for cert in master.certifications:
            line = cert.name
            if cert.issuer:
                line += f", {cert.issuer}"
            if cert.date:
                line += f" ({cert.date})"
            doc.add_paragraph(line, style="List Bullet")

    if master.education:
        doc.add_heading("Education", level=2)
        for edu in master.education:
            line = f"{edu.degree}, {edu.institution}"
            if edu.year:
                line += f" ({edu.year})"
            doc.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def render_resume_md(
    master: MasterResume,
    bullets_by_source: BulletsBySource,
    path: Path,
    include_summary: bool = True,
) -> None:
    lines = [f"# {master.contact.name or 'Resume'}"]
    contact_line = _contact_line(master)
    if contact_line:
        lines.append(contact_line)

    if include_summary and master.summary:
        lines += ["", "## Summary", master.summary]

    lines += ["", "## Experience"]
    for exp in master.experience:
        pairs = bullets_by_source.get(f"experience:{exp.company}")
        bullets = [final for _, final in pairs] if pairs else list(exp.bullets)

        meta = f"{exp.start} - {exp.end}" if (exp.start or exp.end) else ""
        if exp.location:
            meta = f"{meta} | {exp.location}" if meta else exp.location
        lines.append(f"**{exp.role}, {exp.company}**" + (f" ({meta})" if meta else ""))
        lines += [f"- {bullet}" for bullet in bullets]

    project_entries = [
        (proj, bullets_by_source[f"project:{proj.name}"])
        for proj in master.projects
        if f"project:{proj.name}" in bullets_by_source
    ]
    if project_entries:
        lines += ["", "## Projects"]
        for proj, pairs in project_entries:
            lines.append(f"**{proj.name}**")
            lines += [f"- {final}" for _, final in pairs]

    if master.skills:
        lines += ["", "## Skills", ", ".join(master.skills)]

    if master.certifications:
        lines += ["", "## Certifications"]
        for cert in master.certifications:
            line = cert.name
            if cert.issuer:
                line += f", {cert.issuer}"
            if cert.date:
                line += f" ({cert.date})"
            lines.append(f"- {line}")

    if master.education:
        lines += ["", "## Education"]
        for edu in master.education:
            line = f"{edu.degree}, {edu.institution}"
            if edu.year:
                line += f" ({edu.year})"
            lines.append(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
